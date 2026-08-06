from __future__ import annotations

import base64
from datetime import datetime
from html import escape
import json
import os
import re
import tempfile
import unicodedata
import urllib.request
from io import BytesIO
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from PIL import Image, ImageDraw, ImageFont

import cnpj_rf_local


BASE = Path(__file__).resolve().parent
DATA_JSON_CANDIDATES = [
    BASE / "resultados" / "dados_formulario_avaliacao_ex_ante_v18.json",
    BASE / "resultados" / "dados_formulario_avaliacao_ex_ante_v17.json",
    BASE / "resultados" / "dados_formulario_avaliacao_ex_ante_v16.json",
]
RELATORIOS_DIR = BASE / "relatorios_ia"
HOST = "127.0.0.1"
PORT = 8771


def load_payload() -> dict[str, Any]:
    for data_json in DATA_JSON_CANDIDATES:
        if data_json.exists():
            return json.loads(data_json.read_text(encoding="utf-8"))
    return {}


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def strip_accents(text: str) -> str:
    return "".join(ch for ch in unicodedata.normalize("NFD", text) if unicodedata.category(ch) != "Mn")


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        if suffix == ".pdf":
            return normalize_text(extract_pdf(tmp_path))
        if suffix in {".docx", ".doc"}:
            return normalize_text(extract_docx(tmp_path))
        return normalize_text(data.decode("utf-8", errors="ignore"))
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass


def snippet(text: str, start: int, end: int, radius: int = 140) -> str:
    a = max(0, start - radius)
    b = min(len(text), end + radius)
    s = text[a:b].replace("\n", " ")
    return re.sub(r"\s+", " ", s).strip()


def money_to_float(raw: str) -> float | None:
    original = strip_accents(raw).lower()
    multiplier = 1.0
    if re.search(r"\bmilh(?:ao|oes)\b", original):
        multiplier = 1_000_000.0
    elif re.search(r"\bbilh(?:ao|oes)\b", original):
        multiplier = 1_000_000_000.0
    elif re.search(r"\bmil\b", original):
        multiplier = 1_000.0
    raw = re.sub(r"[^\d,\.]", "", raw)
    raw = raw.strip(".,")
    if not raw:
        return None
    if "," in raw and "." in raw:
        raw = raw.replace(".", "").replace(",", ".")
    elif "," in raw:
        raw = raw.replace(",", ".")
    elif "." in raw:
        parts = raw.split(".")
        if len(parts) > 2 or len(parts[-1]) == 3:
            raw = raw.replace(".", "")
    try:
        return float(raw) * multiplier
    except ValueError:
        return None


def format_brl(v: float | None) -> str:
    if v is None:
        return ""
    return f"{int(round(v)):,.0f}".replace(",", ".")


def first_match(text: str, patterns: list[str]) -> tuple[str, str] | None:
    text = re.sub(r"\s+", " ", text)
    for pat in patterns:
        m = re.search(pat, text, flags=re.I | re.S)
        if m:
            value = m.group(1).strip()
            return value, snippet(text, m.start(), m.end())
    return None


def add_field(fields: dict[str, Any], field_id: str, label: str, value: Any, evidence: str, confidence: float, note: str = "") -> None:
    if value in (None, ""):
        return
    fields[field_id] = {
        "label": label,
        "value": str(value),
        "evidence": evidence,
        "confidence": round(float(confidence), 2),
        "note": note,
    }


def find_municipio(text: str, payload: dict[str, Any]) -> tuple[str, str] | None:
    low = strip_accents(text).lower()
    hits = []
    for m in payload.get("municipios", []):
        name = str(m.get("nome", ""))
        if len(name) < 4:
            continue
        idx = low.find(strip_accents(name).lower())
        if idx >= 0:
            hits.append((idx, str(m.get("codigo", "")), name))
    if not hits:
        return None
    idx, code, name = sorted(hits)[0]
    return code, snippet(text, idx, idx + len(name))


def classify_adicionalidade(text: str) -> tuple[str, str, float] | None:
    checks = [
        (r"(seria realizado[^.\n]{0,120}sem o benef[ií]cio|independente do benef[ií]cio|com ou sem o incentivo)", "sim", 0.75),
        (r"(n[aã]o seria realizado[^.\n]{0,120}sem o benef[ií]cio|depende do benef[ií]cio|sem o incentivo[^.\n]{0,80}invi[aá]vel)", "nao", 0.78),
        (r"(menor escala|escala reduzida|parcialmente realizado)", "menor_escala", 0.7),
        (r"(outro estado|outra unidade da federa[cç][aã]o|outra uf)", "outro_estado", 0.68),
    ]
    for pat, value, conf in checks:
        m = re.search(pat, text, flags=re.I)
        if m:
            return value, snippet(text, m.start(), m.end()), conf
    return None


def detect_select_value(text: str, patterns: list[tuple[str, str, float]]) -> tuple[str, str, float] | None:
    for pat, value, conf in patterns:
        m = re.search(pat, text, flags=re.I | re.S)
        if m:
            return value, snippet(text, m.start(), m.end()), conf
    return None


def add_money_field(
    fields: dict[str, Any],
    text: str,
    field_id: str,
    label: str,
    patterns: list[str],
    confidence: float,
) -> None:
    hit = first_match(text, patterns)
    if hit:
        add_field(fields, field_id, label, format_brl(money_to_float(hit[0])), hit[1], confidence)


def add_money_field_if_missing(
    fields: dict[str, Any],
    text: str,
    field_id: str,
    label: str,
    patterns: list[str],
    confidence: float,
) -> None:
    if field_id not in fields:
        add_money_field(fields, text, field_id, label, patterns, confidence)


def add_production_from_two_column_table(fields: dict[str, Any], text: str) -> None:
    compact = re.sub(r"\s+", " ", text)
    patterns = [
        r"(?:valor anual considerado|produ[cç][aã]o anual considerada|produ[cç][aã]o/faturamento)\s*\|\s*(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)\s*(?:por ano)?\s*\|\s*(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
        r"(?:sem benef[ií]cio)\s*\|\s*(?:com benef[ií]cio)\s*\|.{0,120}?(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)\s*\|\s*(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
    ]
    for pat in patterns:
        m = re.search(pat, compact, flags=re.I | re.S)
        if not m:
            continue
        evidence = snippet(compact, m.start(), m.end())
        if "valor_sem_beneficio" not in fields:
            add_field(fields, "valor_sem_beneficio", "Producao/faturamento esperado sem beneficio", format_brl(money_to_float(m.group(1))), evidence, 0.72)
        if "valor" not in fields:
            add_field(fields, "valor", "Producao/faturamento esperado com beneficio", format_brl(money_to_float(m.group(2))), evidence, 0.72)
        return


def extract_fields(text: str, payload: dict[str, Any]) -> dict[str, Any]:
    fields: dict[str, Any] = {}
    warnings: list[str] = []

    hit = first_match(text, [r"\bCNPJ(?:\s*fict[ií]cio)?\D{0,30}(\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2})\b"])
    if hit:
        add_field(fields, "cnpj", "CNPJ", hit[0], hit[1], 0.82, "Campo de caracterizacao. Nao altera os calculos da MIP.")

    hit = first_match(text, [r"\b(?:protocolo|processo|n[uú]mero do processo)\D{0,40}([A-Za-z0-9./-]{4,40})"])
    if hit:
        add_field(fields, "protocolo", "Numero de protocolo", hit[0], hit[1], 0.62)

    hit = first_match(text, [r"\bCNAE(?:\s*(?:principal|fiscal))?\D{0,30}(\d{4,7})\b"])
    if hit:
        add_field(fields, "cnae", "CNAE", hit[0][:4], hit[1], 0.82)

    hit = first_match(text, [r"\b(?:TRU|SCN)\D{0,30}(\d{4})\b"])
    if hit:
        add_field(fields, "tru", "Codigo TRU/SCN", hit[0], hit[1], 0.78)

    ncm_values = sorted(set(re.findall(r"\b\d{8}\b", text)))
    if ncm_values:
        add_field(fields, "ncm", "NCM dos principais produtos", ", ".join(ncm_values[:8]), "NCM(s) de 8 digitos localizados no documento.", 0.72)

    mun = find_municipio(text, payload)
    if mun:
        add_field(fields, "municipio", "Municipio de instalacao", mun[0], mun[1], 0.74, "Codigo IBGE municipal usado pelo formulario.")

    detected = detect_select_value(text, [
        (r"\b(com[eé]rcio|atividade comercial|atacadista|varejista)\b", "comercio", 0.64),
        (r"\b(ind[uú]stria|industrial|fabrica[cç][aã]o|unidade fabril|bebidas|transforma[cç][aã]o)\b", "industria", 0.72),
    ])
    if detected:
        add_field(fields, "macrossegmento", "Macrossegmento", detected[0], detected[1], detected[2])

    detected = detect_select_value(text, [
        (r"(empresa j[aá] existente|reten[cç][aã]o|manuten[cç][aã]o da unidade|risco de sa[ií]da)", "retencao", 0.68),
        (r"(implanta[cç][aã]o nova|nova unidade|nova empresa|expans[aã]o)", "nova", 0.7),
    ])
    if detected:
        add_field(fields, "tipo_analise", "Tipo de analise", detected[0], detected[1], detected[2])

    hit = first_match(text, [
        r"raz[aã]o social\s*\|\s*(.{8,120}?)(?=\s+(?:Nome fantasia|CNPJ|Porte)\s*\|)",
        r"(?:raz[aã]o social|empresa requerente)\s*[:\-]\s*([A-Za-zÀ-ÿ0-9 .,&/-]{8,120})",
    ])
    if hit:
        add_field(fields, "razao_social", "Razao social", hit[0].strip(" .;:-|"), hit[1], 0.66)

    hit = first_match(text, [
        r"nome fantasia\s*\|\s*(.{3,80}?)(?=\s+(?:CNPJ|Porte|UF de origem)\s*\|)",
        r"nome fantasia\s*[:\-]\s*([A-Za-zÀ-ÿ0-9 .,&/-]{3,80})",
    ])
    if hit:
        add_field(fields, "nome_fantasia", "Nome fantasia", hit[0].strip(" .;:-|"), hit[1], 0.64)

    detected = detect_select_value(text, [
        (r"\bMEI\b", "mei", 0.68),
        (r"\bmicroempresa\b", "micro", 0.68),
        (r"\bpequena empresa\b", "pequena", 0.68),
        (r"\bm[eé]dia empresa\b", "media", 0.68),
        (r"\bgrande empresa\b", "grande", 0.7),
    ])
    if detected:
        add_field(fields, "porte_empresa", "Porte", detected[0], detected[1], detected[2])

    uf_hit = first_match(text, [r"UF de origem\s*\|\s*(.{2,90}?)(?=\s+(?:Situa[cç][aã]o|CNAE|C[oó]digo TRU)\s*\|)"])
    uf_scope = uf_hit[0] if uf_hit else text
    detected = detect_select_value(uf_scope, [
        (r"outro estado do Nordeste|Nordeste", "outro_NE", 0.72 if uf_hit else 0.62),
        (r"outro estado fora do Nordeste|fora do Nordeste|Sudeste|Sul|Centro-Oeste", "outro_BR", 0.72 if uf_hit else 0.62),
        (r"exterior|empresa estrangeira|capital estrangeiro", "exterior", 0.72 if uf_hit else 0.62),
        (r"\bPB\b|Para[ií]ba|empresa local", "PB", 0.7 if uf_hit else 0.62),
    ])
    if detected:
        add_field(fields, "uf_origem", "UF de origem", detected[0], uf_hit[1] if uf_hit else detected[1], detected[2])

    detected = detect_select_value(text, [
        (r"implanta[cç][aã]o nova|nova unidade", "nova", 0.7),
        (r"expans[aã]o", "expansao", 0.68),
        (r"reativa[cç][aã]o", "reativacao", 0.68),
        (r"relocaliza[cç][aã]o|transfer[eê]ncia de outra UF", "reloc", 0.68),
    ])
    if detected:
        add_field(fields, "situacao_cadastral", "Situacao", detected[0], detected[1], detected[2])

    add_money_field(fields, text, "valor_sem_beneficio", "Producao/faturamento esperado sem beneficio", [
        r"(?:produ[cç][aã]o/faturamento esperado sem benef[ií]cio|produ[cç][aã]o esperada sem benef[ií]cio|faturamento esperado sem benef[ií]cio|valor sem benef[ií]cio)\D{0,120}(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
        r"(?:sem o benef[ií]cio|sem benef[ií]cio)\D{0,80}(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
    ], 0.78)

    add_money_field(fields, text, "valor", "Producao/faturamento esperado com beneficio", [
        r"(?:produ[cç][aã]o esperada com benef[ií]cio|produ[cç][aã]o/faturamento com benef[ií]cio|faturamento esperado com benef[ií]cio|valor com benef[ií]cio|escala plena)\D{0,120}(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
        r"(?:produ[cç][aã]o anual esperada|valor da produ[cç][aã]o(?: anual)?|faturamento(?: anual)?|receita bruta)\D{0,100}(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
    ], 0.76)

    add_production_from_two_column_table(fields, text)

    if "valor" not in fields:
        add_money_field(fields, text, "valor", "Valor da producao ou faturamento", [
            r"(?:valor do projeto|valor anual)\D{0,100}(?:R\$\s*)?([\d\.\,]+(?:\s*(?:mil|milh(?:[oõ]es|[aã]o)))?)",
        ], 0.6)

    hit = first_match(text, [r"(?:ren[uú]ncia|benef[ií]cio fiscal|cr[eé]dito presumido|redu[cç][aã]o do ICMS)\D{0,80}(\d{1,3}(?:[,.]\d{1,2})?)\s*%"])
    if hit:
        add_field(fields, "renuncia_pct", "Percentual de renuncia do ICMS", hit[0].replace(",", "."), hit[1], 0.76)

    hit = first_match(text, [r"(?:meta de recupera[cç][aã]o de tributos|recupera[cç][aã]o tribut[aá]ria m[ií]nima|meta fiscal)\D{0,80}(\d{1,3}(?:[,.]\d{1,2})?)\s*%"])
    if hit:
        add_field(fields, "meta_recuperacao_tributos", "Meta de recuperacao de tributos", hit[0].replace(",", "."), hit[1], 0.68)

    hit = first_match(text, [r"(?:empregos diretos|postos de trabalho|vagas diretas|empregados diretos)\D{0,80}(\d{1,6})"])
    if hit:
        add_field(fields, "empregos", "Empregos diretos informados", hit[0], hit[1], 0.74)

    hit = first_match(text, [r"(?:sal[aá]rio m[eé]dio(?: mensal)?|remunera[cç][aã]o m[eé]dia)\D{0,80}(?:R\$\s*)?([\d\.\,]{3,})"])
    if hit:
        add_field(fields, "salario", "Salario medio mensal esperado", format_brl(money_to_float(hit[0])), hit[1], 0.62)

    hit = first_match(text, [r"(?:insumos locais|compras locais|fornecedores locais|insumos produzidos localmente)\D{0,80}(\d{1,3}(?:[,.]\d{1,2})?)\s*%"])
    if hit:
        add_field(fields, "local", "Parcela dos insumos produzida localmente (%)", hit[0].replace(",", "."), hit[1], 0.68)

    detected = detect_select_value(text, [
        (r"n[aã]o substitui(?:r[aá]?)? importa[cç][oõ]es", "nao", 0.6),
        (r"substitui(?:r[aá]?)? importa[cç][oõ]es|redu[cç][aã]o de importa[cç][oõ]es", "sim", 0.6),
    ])
    if detected:
        add_field(fields, "substitui", "Produto substitui importacoes", detected[0], detected[1], detected[2])

    detected = detect_select_value(text, [
        (r"produto j[aá] produzido no estado|sem novidade produtiva", "nao", 0.58),
        (r"produto novo|pouco produzido no estado|nova capacidade produtiva", "sim", 0.6),
    ])
    if detected:
        add_field(fields, "novo_produto", "Produto novo ou pouco produzido no estado", detected[0], detected[1], detected[2])

    detected = detect_select_value(text, [
        (r"n[aã]o estrat[eé]gico|sem enquadramento estrat[eé]gico", "nao", 0.58),
        (r"setor estrat[eé]gico|pol[ií]tica industrial|cadeia estrat[eé]gica", "sim", 0.6),
    ])
    if detected:
        add_field(fields, "estrategico", "Setor estrategico para a politica estadual", detected[0], detected[1], detected[2])

    hit = first_match(text, [r"(?:prazo|perman[eê]ncia|tempo do projeto|vig[eê]ncia)\D{0,80}(\d{1,2})\s*(?:anos|ano)"])
    if hit:
        add_field(fields, "permanencia_anos", "Tempo previsto do projeto na Paraiba", hit[0], hit[1], 0.7)

    hit = first_match(text, [r"(?:investimento privado|investimento pr[oó]prio|aporte privado)\D{0,80}(?:R\$\s*)?([\d\.\,]{5,})"])
    if hit:
        add_field(fields, "investimento_privado", "Investimento privado inicial estimado", format_brl(money_to_float(hit[0])), hit[1], 0.7)

    hit = first_match(text, [r"(?:investimento p[uú]blico|contrapartida p[uú]blica|infraestrutura p[uú]blica)\D{0,80}(?:R\$\s*)?([\d\.\,]{5,})"])
    if hit:
        add_field(fields, "investimento_publico", "Investimento publico inicial associado", format_brl(money_to_float(hit[0])), hit[1], 0.64)

    add_money_field(fields, text, "investimento_terreno_imovel", "Valor para aquisicao de terreno ou imovel", [
        r"(?:aquisi[cç][aã]o de terreno ou im[oó]vel|aquisi[cç][aã]o do terreno|valor do terreno|valor do im[oó]vel)\D{0,100}(?:R\$\s*)?([\d\.\,]{5,})",
    ], 0.68)

    add_money_field(fields, text, "investimento_obras", "Valor previsto para obras", [
        r"(?:valor previsto para obras|obras civis|constru[cç][aã]o civil|obras)\D{0,100}(?:R\$\s*)?([\d\.\,]{5,})",
    ], 0.7)

    add_money_field(fields, text, "investimento_outros", "Outros investimentos previstos", [
        r"(?:outros investimentos previstos|P&D|pesquisa e desenvolvimento|treinamentos|capacita[cç][aã]o)\D{0,120}(?:R\$\s*)?([\d\.\,]{5,})",
    ], 0.62)

    hit = first_match(text, [r"(?:ativos\s+(?:fixos\s+)?recuper[aá]veis|parcela\s+dos\s+ativos\s+fixos\s+recuper[aá]veis)\D{0,100}(\d{1,3}(?:[,.]\d{1,2})?)\s*%"])
    if hit:
        add_field(fields, "ativos_recuperaveis_pct", "Parcela dos ativos fixos recuperaveis (%)", hit[0].replace(",", "."), hit[1], 0.62)

    hit = first_match(text, [r"(?:equipamentos adquiridos|m[aá]quinas e equipamentos|parcela dos equipamentos)\D{0,100}(\d{1,3}(?:[,.]\d{1,2})?)\s*%"])
    if hit:
        add_field(fields, "equipamentos_adquiridos_pct", "Parcela dos equipamentos adquiridos para a empresa", hit[0].replace(",", "."), hit[1], 0.62)

    detected = detect_select_value(text, [
        (r"im[oó]vel[^.\n]{0,80}pr[oó]prio|instala[cç][aã]o[^.\n]{0,80}pr[oó]pria", "proprio", 0.68),
        (r"im[oó]vel[^.\n]{0,80}alugado|loca[cç][aã]o do im[oó]vel|galp[aã]o alugado", "alugado", 0.68),
    ])
    if detected:
        add_field(fields, "imovel_tipo", "Imovel de instalacao", detected[0], detected[1], detected[2])

    detected = detect_select_value(text, [
        (r"n[aã]o obteve[^.\n]{0,80}(incentivo locacional|cr[eé]dito|financiamento)", "nao", 0.68),
        (r"obteve[^.\n]{0,80}(incentivo locacional|cr[eé]dito|financiamento)|incentivo locacional", "sim", 0.62),
    ])
    if detected:
        add_field(fields, "incentivo_locacional", "Incentivo locacional ou credito associado", detected[0], detected[1], detected[2])

    ad = classify_adicionalidade(text)
    if ad:
        add_field(fields, "adicionalidade", "Ganho adicional para o Estado", ad[0], ad[1], ad[2])

    detected = detect_select_value(text, [
        (r"mercado local|mercado paraibano|Para[ií]ba", "Mercado local", 0.58),
        (r"fora do estado|outros estados|mercado nacional", "Fora do estado", 0.58),
        (r"exporta[cç][aã]o|mercado externo", "Exportação", 0.58),
        (r"misto|parte[^.\n]{0,80}Para[ií]ba[^.\n]{0,80}fora", "Misto", 0.58),
    ])
    if detected:
        add_field(fields, "destino", "Destino da producao", detected[0], detected[1], detected[2])

    product = None
    product_lines = [
        p.strip(" .;:-|")
        for p in re.findall(r"^([A-Za-zÀ-ÿ0-9 ,;\/-]{8,120})\s*\|\s*\d{8}\s*\|", text, flags=re.I | re.M)
        if p.strip().lower() not in {"produto"}
    ]
    if product_lines:
        add_field(fields, "produtos", "Produtos produzidos/comercializados", "; ".join(product_lines[:6]), "Produtos identificados em tabela com NCM.", 0.68)
    else:
        product = first_match(text, [
            r"produtos?\s+(?:produzidos|declarados|no estabelecimento incentivado)?:?\s*([A-Za-zÀ-ÿ0-9 ,;\/-]{12,180})",
            r"produzir[aá]?\s+e\s+comercializar[aá]?\D{0,80}:?\s*([A-Za-zÀ-ÿ0-9 ,;\/-]{12,180})",
            r"produtos?\D{0,40}([A-Za-zÀ-ÿ0-9 ,;\/-]{12,160})",
        ])
    if "produtos" not in fields and product:
        add_field(fields, "produtos", "Produtos produzidos/comercializados", product[0].strip(" .;:"), product[1], 0.48)

    if not fields:
        warnings.append("Nao foram encontrados campos com confianca minima. Verifique se o documento contem texto pesquisavel ou use OCR antes da leitura.")
    if "valor" not in fields:
        warnings.append("Valor da producao/faturamento nao foi identificado automaticamente.")
    if "renuncia_pct" not in fields:
        warnings.append("Percentual de renuncia nao foi identificado automaticamente.")

    return {
        "fields": fields,
        "warnings": warnings,
        "text_preview": text[:1200],
        "engine": "extracao_estruturada_local",
    }


def score_label(score: float | None) -> str:
    if score is None:
        return "sem nota calculada"
    if score < 5:
        return "sinal vermelho, indicando que o beneficio requer mais informacoes"
    if score < 7:
        return "sinal laranja, indicando necessidade de analise complementar"
    return "sinal verde, com aprovacao preliminar condicionada a validacao documental"


def rule_based_analysis(ctx: dict[str, Any]) -> str:
    score = ctx.get("score")
    try:
        score_num = float(score)
    except Exception:
        score_num = None
    sig = score_label(score_num)
    setor = ctx.get("setor") or "setor informado"
    municipio = ctx.get("municipio") or "municipio nao informado"
    parts = [
        "Sintese preliminar\n"
        f"A avaliacao preliminar do empreendimento no setor {setor}, no municipio {municipio}, resultou em {sig}. "
        "A nota deve ser lida como uma triagem tecnica: ela organiza evidencias, mas nao substitui a verificacao documental nem a decisao administrativa.",
    ]
    risk = ctx.get("rentRisk", {})
    if risk:
        parts.append(
            "Risco de rent-seeking\n"
            f"O risco de rent-seeking foi classificado como {risk.get('level', '-')}. Esse card indica a possibilidade de o beneficio gerar apropriacao privada sem contrapartida suficiente, especialmente quando ha baixo ganho adicional para o Estado, compras locais reduzidas, poucos empregos ou facilidade de saida do estado. "
            f"Alertas registrados: {'; '.join(risk.get('alerts') or ['sem alerta relevante informado'])}."
        )
    fiscal = ctx.get("taxRenuncia")
    if fiscal is not None:
        parts.append(
            "Retorno fiscal e tributos\n"
            f"No retorno fiscal, a razao entre tributos liquidos estimados e renuncia foi de {float(fiscal):.2f}. Valores abaixo de 1 indicam que a recuperacao fiscal estrita e inferior ao custo da renuncia no ano-base; valores acima de 1 sugerem recuperacao fiscal mais favoravel. "
            "Quando ha renuncia parcial, o painel considera a parcela direta nao renunciada e os tributos indiretos estimados pelos encadeamentos produtivos."
        )
    social = ctx.get("socialVaRenuncia")
    if social is not None:
        parts.append(
            "Viabilidade social pelo valor adicionado\n"
            f"Na viabilidade social, a relacao entre valor adicionado em valor presente e beneficio tributario em valor presente foi de {float(social):.2f}. Esse indicador nao mede caixa do governo; ele mede o beneficio economico local comparado ao custo fiscal. "
            "Ele ajuda a separar projetos que apenas deslocam faturamento daqueles que ampliam renda gerada internamente."
        )
    wage = ctx.get("wageRenuncia")
    if wage is not None:
        parts.append(
            "Massa salarial\n"
            f"Na analise de custo-beneficio salarial, a razao entre massa salarial estimada e renuncia em valor presente foi de {float(wage):.2f}. Essa leitura e util para auditores porque traduz parte do beneficio economico em renda do trabalho, uma dimensao mais intuitiva que o multiplicador agregado."
        )
    jobs = ctx.get("totalJobs")
    if jobs is not None:
        parts.append(
            "Emprego\n"
            f"O impacto sobre emprego combina empregos diretos declarados e empregos indiretos estimados pela MIP, totalizando aproximadamente {float(jobs):,.0f} postos. ".replace(",", ".")
            + f"Os empregos diretos declarados foram {ctx.get('directJobs', 0)} e os indiretos estimados foram {float(ctx.get('indirectJobs') or 0):,.0f}.".replace(",", ".")
        )
    coverage = ctx.get("territorialCoverage")
    if coverage is not None:
        parts.append(
            "Territorio e especializacao produtiva\n"
            f"A leitura territorial aponta que {coverage} dos municipios apresentam especializacao robusta em ao menos um dos setores indiretamente impactados. Quanto maior essa abrangencia, maior o potencial de espalhamento territorial dos efeitos indiretos, desde que as compras locais sejam efetivas. "
            f"No municipio declarado, {ctx.get('qlTopCount', 0)} dos setores indiretamente impactados aparecem com especializacao robusta."
        )
    local_share = ctx.get("localShare")
    if local_share is not None:
        parts.append(
            "Compras locais e adensamento produtivo\n"
            f"A parcela declarada de insumos produzidos localmente foi de {float(local_share) * 100:.1f}%. Quanto maior essa parcela, maior a chance de o choque de demanda permanecer na economia paraibana em vez de vazar para importacoes de outros estados ou paises."
        )
    if ctx.get("techLabel"):
        parts.append(
            "Conteudo tecnologico\n"
            f"O conteudo tecnologico associado ao setor/produto foi classificado como {ctx.get('techLabel')}, com score setorial {ctx.get('techScore', '-')}. Projetos com maior conteudo tecnologico tendem a receber melhor leitura preliminar quando a politica busca diversificacao produtiva e aumento de complexidade."
        )
    sectors = ctx.get("topIndirectSectors") or []
    if sectors:
        sector_txt = "; ".join(
            f"{s.get('codigo')} - {s.get('setor')} ({float(s.get('participacao') or 0) * 100:.1f}%)"
            for s in sectors[:8]
        )
        parts.append(
            "Setores indiretamente impactados\n"
            "Os principais setores que concentram os impactos indiretos sobre a producao foram: "
            + sector_txt
            + ". Essa lista e importante porque o mapa territorial deve ser lido a partir desses efeitos indiretos, e nao apenas pelo setor diretamente beneficiado."
        )
    missing = ctx.get("missingQualifiers") or []
    if missing:
        parts.append(
            "Lacunas de informacao\n"
            "Campos qualificadores nao informados ou frageis devem ser tratados como risco: " + "; ".join(missing[:8]) + ". Quando esses itens recebem peso maior que zero no painel, a ausencia de informacao reduz a nota preliminar."
        )
    parts.append(
        "Recomendacao operacional\n"
        "Usar este texto como minuta analitica, conferir evidencias documentais, revisar campos extraidos automaticamente e registrar no parecer quais hipoteses foram usadas. A IA funciona como assistente de leitura documental e redacao analitica; os calculos da MIP, a nota e os indicadores permanecem determinados pelo painel."
    )
    return "\n\n".join(parts)


def analysis_pdf_base64(analysis: str, ctx: dict[str, Any]) -> str:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
        title="Analise textual assistida por IA",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SmallMeta", parent=styles["Normal"], fontSize=8.5, leading=11, textColor=colors.HexColor("#475569")))
    styles.add(ParagraphStyle(name="SectionBody", parent=styles["Normal"], fontSize=10.5, leading=14, spaceAfter=7))
    story = [
        Paragraph("Análise textual assistida por IA", styles["Title"]),
        Paragraph(
            f"Setor: {escape(str(ctx.get('codigo', '-')))} - {escape(str(ctx.get('setor', '-')))}<br/>"
            f"Município: {escape(str(ctx.get('municipio', '-')))}<br/>"
            f"Nota preliminar: {escape(str(ctx.get('score', '-')))} | Decisão preliminar: {escape(str(ctx.get('signal', '-')))}<br/>"
            f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            styles["SmallMeta"],
        ),
        Spacer(1, 0.35 * cm),
    ]
    for block in analysis.split("\n\n"):
        clean = block.strip()
        if not clean:
            continue
        lines = clean.split("\n", 1)
        if len(lines) == 2 and len(lines[0]) < 80:
            story.append(Paragraph(f"<b>{escape(lines[0])}</b>", styles["Heading2"]))
            story.append(Paragraph(escape(lines[1]).replace("\n", "<br/>"), styles["SectionBody"]))
        else:
            story.append(Paragraph(escape(clean).replace("\n", "<br/>"), styles["SectionBody"]))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("Observação: a IA funciona como assistente de leitura documental e redação analítica. Os cálculos e a nota são produzidos pelo motor determinístico do painel.", styles["SmallMeta"]))
    doc.build(story)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def maybe_ollama_analysis(ctx: dict[str, Any]) -> str | None:
    model = os.environ.get("OLLAMA_MODEL", "").strip()
    if not model:
        return None
    prompt = (
        "INSTRUCAO CRITICA: voce NAO deve recalcular, corrigir, substituir ou questionar nenhum numero "
        "apresentado no JSON. Os valores foram calculados por motor deterministico externo e sao os dados "
        "oficiais desta avaliacao. Sua funcao e exclusivamente INTERPRETAR, CONTEXTUALIZAR e REDIGIR.\n\n"
        "Escreva em portugues um relatorio tecnico-analitico para auditores fiscais sobre uma avaliacao ex ante "
        "de beneficio tributario. A IA funciona como assistente de leitura documental e redacao analitica; "
        "os calculos e a nota foram gerados por um motor deterministico. Use a chave memoriaCalculo, quando "
        "disponivel, para detalhar a memoria de calculo sem refazer contas. Explique, em linguagem clara, a nota "
        "preliminar, a leitura dos multiplicadores, os riscos, o retorno fiscal, valor adicionado, emprego, "
        "massa salarial e territorio. Separe a resposta em blocos curtos com titulos. "
        "Nao invente numeros, nao altere a nota e use apenas o JSON.\n\nJSON:\n"
        + json.dumps(ctx, ensure_ascii=False)
    )
    data = json.dumps({"model": model, "prompt": prompt, "stream": False}, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request("http://127.0.0.1:11434/api/generate", data=data, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            body = json.loads(resp.read().decode("utf-8"))
        return body.get("response")
    except Exception:
        return None


def generate_analysis(ctx: dict[str, Any]) -> tuple[str, str, str]:
    model = os.environ.get("OLLAMA_MODEL", "").strip()
    ollama_text = maybe_ollama_analysis(ctx)
    if ollama_text:
        return ollama_text, "ollama", model
    return rule_based_analysis(ctx), "analise_regrada_local", ""


def as_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        return float(value)
    except Exception:
        return default


def money(value: Any) -> str:
    v = as_float(value)
    return "R$ " + f"{v:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")


def pct(value: Any, digits: int = 1) -> str:
    return f"{as_float(value):.{digits}f}%".replace(".", ",")


def num(value: Any, digits: int = 1) -> str:
    return f"{as_float(value):,.{digits}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def plain(value: Any, default: str = "-") -> str:
    if value is None or value == "":
        return default
    return str(value)


def hybrid_fallback_narratives(ctx: dict[str, Any]) -> dict[str, str]:
    mods = ctx.get("modulosV7", {}) or {}
    neutral = (ctx.get("neutralidadeFiscalExperimental") or {})
    fiscal_msg = "atendida" if neutral.get("atingida") else "não atendida"
    return {
        "resumo": (
            f"A avaliação preliminar resultou em {plain(ctx.get('signal')).lower()}, com nota {num(ctx.get('score'), 1)} de 10. "
            f"O setor analisado é {plain(ctx.get('codigo'))} - {plain(ctx.get('setor'))}, no município de {plain(ctx.get('municipio'))}. "
            "A leitura deve ser usada como triagem técnica, não como decisão automática."
        ),
        "economia": (
            "Os impactos econômicos foram calculados a partir do choque de produção informado no formulário e dos multiplicadores da MIP. "
            "A interpretação principal é observar se o projeto gera produção, valor adicionado, emprego e massa salarial de forma compatível com a renúncia fiscal."
        ),
        "fiscal": (
            f"O teste de neutralidade fiscal foi {fiscal_msg}. A lógica é comparar os tributos indiretos gerados pelo acréscimo de produção com a meta de recuperação da renúncia definida pelo gestor."
        ),
        "territorio": (
            "A leitura territorial identifica onde os efeitos indiretos podem ser absorvidos dentro da Paraíba. Municípios com maior capacidade territorial combinam especialização produtiva, massa econômica, proximidade e posição na rede urbana."
        ),
        "qualidade": (
            f"A qualidade da informação foi classificada como {plain((mods.get('qualidadeInformacao') or {}).get('level'))}. "
            f"A plausibilidade econômica foi classificada como {plain((mods.get('plausibilidadeEconomica') or {}).get('level'))}, e o ganho adicional para o Estado como {plain((mods.get('adicionalidade') or {}).get('level'))}."
        ),
        "riscos": (
            "Os alertas devem orientar diligências adicionais. Campos ausentes, neutralidade fiscal não atendida, baixa compra local ou alto risco de benefício pouco produtivo não reprovam automaticamente o pleito, mas aumentam a necessidade de comprovação documental."
        ),
        "recomendacao": (
            "A recomendação operacional é conferir os dados declarados, validar documentos comprobatórios e usar este relatório como apoio técnico ao parecer fiscal e à decisão administrativa competente."
        ),
    }


def maybe_ollama_hybrid_narratives(ctx: dict[str, Any]) -> tuple[dict[str, str], str, str]:
    model = os.environ.get("OLLAMA_MODEL", "").strip()
    if not model:
        return hybrid_fallback_narratives(ctx), "analise_regrada_local", ""
    compact = {
        "empresa": ctx.get("empresa"),
        "tipoAnalise": ctx.get("tipoAnalise"),
        "macrossegmento": ctx.get("macrossegmento"),
        "codigo": ctx.get("codigo"),
        "setor": ctx.get("setor"),
        "municipio": ctx.get("municipio"),
        "score": ctx.get("score"),
        "signal": ctx.get("signal"),
        "impactosEsperados": ctx.get("impactosEsperados"),
        "neutralidadeFiscalExperimental": ctx.get("neutralidadeFiscalExperimental"),
        "rentRisk": ctx.get("rentRisk"),
        "modulosV7": ctx.get("modulosV7"),
        "territorialAbsorption": ctx.get("territorialAbsorption"),
        "topIndirectSectors": ctx.get("topIndirectSectors"),
        "missingQualifiers": ctx.get("missingQualifiers"),
        "memoriaCalculo": ctx.get("memoriaCalculo"),
    }
    prompt = (
        "Voce é assistente de redação analítica para auditores fiscais. "
        "Responda SOMENTE um JSON válido, sem markdown, com as chaves: "
        "resumo, economia, fiscal, territorio, qualidade, riscos, recomendacao. "
        "Nao crie nem altere numeros. Nao escreva tabelas. Nao mude titulos. "
        "A estrutura, tabelas, numeros, alertas e ressalvas serão controlados pelo painel. "
        "Sua função é apenas redigir interpretação clara, conectar indicadores e explicar a leitura para auditor com pouca familiaridade com MIP. "
        "Use no maximo 120 palavras por chave.\n\nJSON de insumos:\n"
        + json.dumps(compact, ensure_ascii=False)
    )
    data = json.dumps({"model": model, "prompt": prompt, "stream": False}, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request("http://127.0.0.1:11434/api/generate", data=data, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=45) as resp:
            body = json.loads(resp.read().decode("utf-8"))
        raw = (body.get("response") or "").strip()
        parsed = json.loads(raw)
        if isinstance(parsed, dict):
            base = hybrid_fallback_narratives(ctx)
            base.update({k: str(v) for k, v in parsed.items() if k in base and v})
            return base, "ollama", model
    except Exception:
        pass
    return hybrid_fallback_narratives(ctx), "analise_regrada_local", ""


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(14, 79, 148)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(15, 23, 42)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(71, 85, 105)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = plain(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, Any]]) -> None:
    add_table(doc, ["Indicador", "Resultado"], [[a, b] for a, b in rows])


def chart_png_base64(title: str, rows: list[tuple[str, float]], width: int = 1100, height: int = 520) -> str | None:
    rows = [(str(label)[:64], float(value or 0)) for label, value in rows if float(value or 0) > 0][:10]
    if not rows:
        return None
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("Arial.ttf", 26)
        font = ImageFont.truetype("Arial.ttf", 18)
    except Exception:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()
    draw.text((28, 24), title, fill=(15, 23, 42), font=font_title)
    max_v = max(v for _, v in rows) or 1
    top = 78
    bar_h = 30
    gap = 14
    left = 360
    right = width - 70
    for idx, (label, value) in enumerate(rows):
        y = top + idx * (bar_h + gap)
        draw.text((28, y + 4), label, fill=(51, 65, 85), font=font)
        bar_w = int((right - left) * value / max_v)
        draw.rectangle((left, y, left + bar_w, y + bar_h), fill=(14, 79, 148))
        draw.text((left + bar_w + 10, y + 5), num(value, 1), fill=(15, 23, 42), font=font)
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def add_png_b64(doc: Document, png_b64: str | None, width_inches: float = 6.6) -> None:
    if not png_b64:
        return
    try:
        if "," in png_b64 and png_b64.strip().startswith("data:"):
            png_b64 = png_b64.split(",", 1)[1]
        data = base64.b64decode(png_b64)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(data)
            path = tmp.name
        doc.add_picture(path, width=Inches(width_inches))
        Path(path).unlink(missing_ok=True)
    except Exception:
        add_note(doc, "A imagem não pôde ser incorporada automaticamente ao Word.")


def hybrid_docx_bytes(ctx: dict[str, Any], map_png_base64: str | None = None) -> tuple[bytes, str, str]:
    narratives, engine, model = maybe_ollama_hybrid_narratives(ctx)
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Relatório híbrido de avaliação ex ante de benefício tributário")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(14, 79, 148)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("SEFAZ-PB | Matriz Insumo-Produto da Paraíba | Produto experimental").italic = True
    doc.add_paragraph()

    empresa = ctx.get("empresa") or {}
    doc.add_heading("1. Identificação e sinal preliminar", level=1)
    add_kv_table(doc, [
        ("Empresa", empresa.get("razao_social") or "-"),
        ("CNPJ", empresa.get("cnpj") or "-"),
        ("Protocolo", ctx.get("protocolo") or "-"),
        ("Tipo de análise", ctx.get("tipoAnalise") or "-"),
        ("Macrossegmento", ctx.get("macrossegmento") or "-"),
        ("Setor MIP/TRU", f"{plain(ctx.get('codigo'))} - {plain(ctx.get('setor'))}"),
        ("Município", ctx.get("municipio") or "-"),
        ("Nota preliminar", f"{num(ctx.get('score'), 1)} de 10"),
        ("Sinal preliminar", ctx.get("signal") or "-"),
        ("Motor de redação", f"{engine}{' | ' + model if model else ''}"),
        ("Data de geração", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ])

    cadastro_rf = empresa.get("cadastro_rf") or {}
    if cadastro_rf:
        endereco = cadastro_rf.get("endereco") or {}
        endereco_texto = ", ".join(
            plain(value)
            for value in [
                " ".join(filter(None, [endereco.get("tipo_logradouro"), endereco.get("logradouro")])),
                endereco.get("numero"),
                endereco.get("complemento"),
                endereco.get("bairro"),
                " / ".join(filter(None, [endereco.get("municipio"), endereco.get("uf")])),
                endereco.get("cep"),
            ]
            if value
        ) or "-"
        cnae_principal = cadastro_rf.get("cnae_principal") or {}
        doc.add_heading("1.1 Caracterização cadastral na base da Receita Federal", level=2)
        add_kv_table(doc, [
            ("Situação cadastral", cadastro_rf.get("situacao_cadastral") or "-"),
            ("Matriz/filial", cadastro_rf.get("matriz_filial") or "-"),
            ("Início da atividade", cadastro_rf.get("data_inicio_atividade") or "-"),
            ("Porte", cadastro_rf.get("porte") or "-"),
            ("Natureza jurídica", cadastro_rf.get("natureza_juridica") or "-"),
            ("Capital social", money(cadastro_rf.get("capital_social"))),
            ("Endereço cadastral", endereco_texto),
            ("CNAE principal", f"{plain(cnae_principal.get('codigo'))} - {plain(cnae_principal.get('descricao'))}"),
            ("Optante do Simples Nacional", cadastro_rf.get("simples_nacional") or "-"),
            ("MEI", cadastro_rf.get("mei") or "-"),
            ("Competência da base", cadastro_rf.get("competencia") or "-"),
        ])
        activities = cadastro_rf.get("atividades_secundarias") or []
        if activities:
            add_table(doc, ["CNAE secundário", "Descrição"], [
                [item.get("codigo") or "-", item.get("descricao") or "-"]
                for item in activities[:12]
            ])
        add_note(
            doc,
            "Fonte: Receita Federal do Brasil, Dados Abertos do CNPJ, consultada em índice local. "
            "Os dados cadastrais não substituem certidões, diligência fiscal ou comprovação documental.",
        )
    else:
        add_note(doc, "A caracterização cadastral da Receita Federal não foi incorporada a esta análise.")

    doc.add_heading("2. Resumo executivo para decisão", level=1)
    doc.add_paragraph(narratives["resumo"])
    resumo = ((ctx.get("modulosV7") or {}).get("resumoDecisorio") or {})
    add_table(doc, ["Pontos favoráveis", "Pontos de atenção"], [[
        "\n".join(resumo.get("positives") or ["Sem ponto favorável automático destacado."]),
        "\n".join(resumo.get("cautions") or ["Sem alerta automático adicional."]),
    ]])

    memoria = ctx.get("memoriaCalculo") or {}
    choque = memoria.get("choque") or {}
    impactos = ctx.get("impactosEsperados") or {}
    doc.add_heading("3. Dados usados e choque aplicado à MIP", level=1)
    add_kv_table(doc, [
        ("Valor com benefício informado", money(choque.get("valorComBeneficioDeclarado") or ctx.get("declaredValue"))),
        ("Valor sem benefício informado", money(choque.get("valorSemBeneficioDeclarado") or ctx.get("valorSemBeneficioDeclarado"))),
        ("Choque considerado na MIP", money(choque.get("valorIncrementalConsideradoNaMip") or ctx.get("valueBRL"))),
        ("Margem comercial aplicada", "-" if choque.get("margemComercial") is None else pct(choque.get("margemComercial"))),
        ("Percentual de renúncia solicitado", pct(ctx.get("renunciaPctSolicitada"), 2)),
        ("Percentual de renúncia considerado", pct(ctx.get("renunciaPct"), 2)),
        ("Meta de recuperação de tributos", pct(ctx.get("metaRecuperacaoTributos"), 1)),
    ])

    doc.add_heading("4. Impactos econômicos estimados", level=1)
    doc.add_paragraph(narratives["economia"])
    add_kv_table(doc, [
        ("Impacto na produção", money(impactos.get("producao"))),
        ("Impacto no valor adicionado", money(impactos.get("valorAdicionado"))),
        ("Empregos na operação", num(impactos.get("empregos"), 1)),
        ("Empregos na implantação por obras", num(impactos.get("empregosImplantacaoTotal"), 1)),
        ("Massa salarial estimada pela MIP", money(impactos.get("massaSalarialMip"))),
        ("Multiplicador de produção", num(ctx.get("producaoMultiplicador"), 4)),
        ("Multiplicador de VA", num(ctx.get("vaMultiplicador"), 4)),
        ("Multiplicador de emprego por R$ 1 milhão", num(ctx.get("empregoMultiplicador"), 4)),
    ])
    sectors = ctx.get("topIndirectSectors") or []
    add_table(doc, ["SCN", "Setor indiretamente impactado", "Participação", "Impacto"], [
        [s.get("codigo"), s.get("setor"), pct(as_float(s.get("participacao")) * 100, 1), money(s.get("impacto"))]
        for s in sectors[:12]
    ] or [["-", "Sem setores indiretos informados.", "-", "-"]])
    chart = chart_png_base64(
        "Setores com maior impacto indireto",
        [(f"{s.get('codigo')} - {s.get('setor')}", as_float(s.get("impacto"))) for s in sectors],
    )
    add_png_b64(doc, chart)

    doc.add_heading("5. Leitura fiscal e neutralidade da renúncia", level=1)
    doc.add_paragraph(narratives["fiscal"])
    neutral = ctx.get("neutralidadeFiscalExperimental") or {}
    trib = memoria.get("tributos") or {}
    add_kv_table(doc, [
        ("Renúncia fiscal estimada", money(trib.get("renuncia"))),
        ("Tributos indiretos estimados", money(trib.get("indiretos"))),
        ("Tributos líquidos estimados", money(trib.get("liquidos"))),
        ("Tributos líquidos / renúncia", "-" if ctx.get("taxRenuncia") is None else num(ctx.get("taxRenuncia"), 2)),
        ("Neutralidade fiscal", "Atendida" if neutral.get("atingida") else "Não atendida"),
        ("Diferença em relação à meta", money(neutral.get("diferencaFiscal"))),
        ("Acréscimo mínimo para atingir a meta", "-" if neutral.get("incrementoNeutro") is None else money(neutral.get("incrementoNeutro"))),
        ("Renúncia máxima compatível com a meta", "-" if neutral.get("renunciaMaximaNeutra") is None else pct(as_float(neutral.get("renunciaMaximaNeutra")) * 100, 2)),
    ])

    doc.add_heading("6. Emprego e massa salarial", level=1)
    add_kv_table(doc, [
        ("Empregos diretos informados/preservados", num(ctx.get("directJobs"), 1)),
        ("Empregos indiretos estimados", num(ctx.get("indirectJobs"), 1)),
        ("Empregos totais na operação", num(ctx.get("totalJobs"), 1)),
        ("Empregos de implantação - diretos nas obras", num(impactos.get("empregosImplantacaoDiretos"), 1)),
        ("Empregos de implantação - indiretos nas obras", num(impactos.get("empregosImplantacaoIndiretos"), 1)),
        ("Empregos totais na implantação", num(impactos.get("empregosImplantacaoTotal"), 1)),
        ("Massa salarial estimada pela MIP", money(impactos.get("massaSalarialMip"))),
        ("Massa salarial / renúncia em VP", "-" if ctx.get("wageRenuncia") is None else num(ctx.get("wageRenuncia"), 2)),
    ])

    doc.add_heading("7. Leitura territorial", level=1)
    doc.add_paragraph(narratives["territorio"])
    mapa = ctx.get("mapaRegic") or {}
    loc = ctx.get("locationalAssessment") or ((ctx.get("indicadoresMunicipais") or {}).get("desconcentracaoEconomica") or {})
    scores = mapa.get("scores") or []
    top_scores = [s for s in scores if s.get("capacidade_alta")][:12]
    add_kv_table(doc, [
        ("Município de origem", ctx.get("municipio") or "-"),
        ("Zona locacional", plain(loc.get("zona") or "-")),
        ("Contribuição para desconcentração econômica", f"{plain(loc.get('level') or '-')} ({num(as_float(loc.get('score')) * 100, 0)}%)"),
        ("Cobertura territorial", ctx.get("territorialCoverage") or "-"),
        ("Municípios com alta capacidade territorial", plain(((ctx.get("territorialAbsorption") or {}).get("count")))),
        ("Setores indiretos considerados no mapa", plain(ctx.get("impactSectorsCount"))),
    ])
    if loc.get("message"):
        add_note(doc, plain(loc.get("message")))
    add_png_b64(doc, map_png_base64, width_inches=6.8)
    add_table(doc, ["Município", "Score", "Capacidade", "Distância km", "Principal setor fornecedor"], [
        [s.get("nome"), num(s.get("score"), 1), s.get("capacidade_classe"), num(s.get("dist_km"), 1), s.get("principal_setor")]
        for s in top_scores
    ] or [["-", "-", "Sem município destacado com alta capacidade.", "-", "-"]])
    chart2 = chart_png_base64(
        "Municípios com maior capacidade territorial",
        [(s.get("nome"), as_float(s.get("score"))) for s in top_scores],
    )
    add_png_b64(doc, chart2)

    doc.add_heading("8. Qualidade da informação, plausibilidade e ganho adicional para o Estado", level=1)
    doc.add_paragraph(narratives["qualidade"])
    mods = ctx.get("modulosV7") or {}
    qual = mods.get("qualidadeInformacao") or {}
    plaus = mods.get("plausibilidadeEconomica") or {}
    adic = mods.get("adicionalidade") or {}
    sens = mods.get("sensibilidade") or {}
    sens_nota = mods.get("sensibilidadeNota") or {}
    add_kv_table(doc, [
        ("Qualidade da informação", f"{plain(qual.get('level'))} ({num(as_float(qual.get('score')) * 100, 0)}%)"),
        ("Plausibilidade econômica", plain(plaus.get("level"))),
        ("Ganho adicional para o Estado", f"{plain(adic.get('level'))} ({num(as_float(adic.get('score')) * 100, 0)}%)"),
        ("Sensibilidade", plain(sens.get("leitura"))),
        ("Campos ausentes/frágeis", "; ".join(ctx.get("missingQualifiers") or qual.get("missing") or []) or "-"),
    ])
    doc.add_heading("8.1 Análise de sensibilidade da nota", level=2)
    add_note(doc, "Esta seção não indica aprovação automática. As combinações consideram apenas dimensões negociáveis ou pactuáveis, como emprego, compras locais, permanência, desenho fiscal e compromissos verificáveis. Características intrínsecas, como setor-chave, conteúdo tecnológico típico do produto e multiplicadores setoriais, aparecem apenas como diagnóstico.")
    add_kv_table(doc, [
        ("Nota atual", num(sens_nota.get("currentScore"), 1)),
        ("Meta de referência", num(sens_nota.get("target") or 7, 1)),
        ("Distância até a meta", num(sens_nota.get("gap"), 1)),
        ("Combinações negociáveis encontradas", plain(len(sens_nota.get("viable") or []))),
    ])
    add_table(doc, ["Atributo", "Tipo", "Situação", "Peso", "Ganho potencial", "Entra nas combinações?", "Leitura"], [
        [l.get("label"), l.get("type"), l.get("status"), num(l.get("weight"), 2), num(l.get("gain"), 2), "Sim" if l.get("actionable") else "Não", l.get("note")]
        for l in (sens_nota.get("levers") or [])[:12]
    ] or [["-", "-", "-", "-", "-", "-", "Sem alavancas calculadas."]])
    combos = (sens_nota.get("viable") or sens_nota.get("near") or [])[:6]
    add_table(doc, ["Condições", "Tipo", "Ganho", "Nota simulada", "Leitura"], [
        [
            "\n".join(plain(i.get("label")) for i in (c.get("items") or [])),
            "\n".join(plain(i.get("type")) for i in (c.get("items") or [])),
            num(c.get("gain"), 2),
            num(c.get("score"), 1),
            "Atinge a meta, se comprovado/pactuado." if as_float(c.get("score")) >= as_float(sens_nota.get("target") or 7) else "Combinação próxima, mas ainda abaixo da meta.",
        ]
        for c in combos
    ] or [["-", "-", "-", "-", "Sem combinação negociável suficiente com os critérios atuais."]])
    doc.add_heading("8.2 Sensibilidade fiscal", level=2)
    scenarios = sens.get("scenarios") or []
    add_table(doc, ["Cenário", "Choque", "Neutralidade", "Diferença fiscal", "Produção", "Empregos"], [
        [s.get("nome"), money(s.get("incremento")), "Atendida" if s.get("neutralidade") else "Não atendida", money(s.get("diferenca")), money(s.get("producao")), num(s.get("empregos"), 1)]
        for s in scenarios
    ] or [["-", "-", "-", "-", "-", "-"]])

    doc.add_heading("9. Riscos, diligências e ressalvas metodológicas", level=1)
    doc.add_paragraph(narratives["riscos"])
    risk = ctx.get("rentRisk") or {}
    add_table(doc, ["Tipo", "Descrição"], (
        [["Risco de benefício pouco produtivo", f"{plain(risk.get('level'))}. Alertas: " + "; ".join(risk.get("alerts") or [])]]
        + [["Campo a verificar", item] for item in (ctx.get("missingQualifiers") or [])]
        + [["Ressalva metodológica", "A IA funciona como assistente de leitura documental e redação analítica. Os cálculos, pesos, indicadores, tabelas e ordem das seções são definidos pelo painel."]]
        + [["Ressalva metodológica", "Os resultados da MIP representam impactos médios setoriais e não substituem diligência fiscal, análise jurídica, comprovação documental ou decisão administrativa."]]
        + [["Ressalva metodológica", "O mapa territorial é uma aproximação baseada em QL, REGIC, distância rodoviária e massa econômica municipal; ele indica potencial de absorção, não garantia de contratação local."]]
    ))

    doc.add_heading("10. Recomendação operacional", level=1)
    doc.add_paragraph(narratives["recomendacao"])

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), engine, model


def json_response(handler: BaseHTTPRequestHandler, status: int, payload: dict[str, Any]) -> None:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Access-Control-Allow-Origin", "*")
    handler.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
    handler.send_header("Access-Control-Allow-Headers", "Content-Type")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


class Handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self) -> None:
        json_response(self, 200, {"ok": True})

    def do_GET(self) -> None:
        path = unquote(urlparse(self.path).path)
        if path == "/health":
            json_response(
                self,
                200,
                {
                    "ok": True,
                    "service": "servico_ia_contratos",
                    "port": PORT,
                    "base_cnpj_rf": {
                        "disponivel": cnpj_rf_local.CNPJ_RF_DB.exists(),
                        "arquivo": str(cnpj_rf_local.CNPJ_RF_DB),
                    },
                },
            )
            return
        if path.startswith("/consulta_cnpj/"):
            cnpj = path.removeprefix("/consulta_cnpj/")
            try:
                result = cnpj_rf_local.consultar_cnpj_base_rf(cnpj)
                json_response(self, 200 if result.get("ok") else 404, result)
            except FileNotFoundError as exc:
                json_response(self, 503, {"ok": False, "error": str(exc)})
            except ValueError as exc:
                json_response(self, 400, {"ok": False, "error": str(exc)})
            return
        json_response(self, 404, {"ok": False, "error": "Endpoint nao encontrado."})

    def do_POST(self) -> None:
        try:
            length = int(self.headers.get("Content-Length", "0"))
            raw = self.rfile.read(length)
            payload = json.loads(raw.decode("utf-8"))
            if self.path.startswith("/extrair_contrato"):
                filename = payload.get("filename", "documento")
                content = base64.b64decode(payload.get("content_base64", ""))
                text = extract_text(filename, content)
                result = extract_fields(text, load_payload())
                json_response(self, 200, {"ok": True, **result})
                return
            if self.path.startswith("/gerar_analise"):
                ctx = payload.get("context", {})
                text, engine, model = generate_analysis(ctx)
                json_response(self, 200, {"ok": True, "analysis": text, "engine": engine, "model": model})
                return
            if self.path.startswith("/gerar_pdf_analise"):
                ctx = payload.get("context", {})
                text, engine, model = generate_analysis(ctx)
                filename_code = re.sub(r"[^A-Za-z0-9_-]+", "_", str(ctx.get("codigo") or "setor")).strip("_") or "setor"
                json_response(
                    self,
                    200,
                    {
                        "ok": True,
                        "analysis": text,
                        "pdf_base64": analysis_pdf_base64(text, ctx),
                        "filename": f"analise_ia_{filename_code}.pdf",
                        "engine": engine,
                        "model": model,
                    },
                )
                return
            if self.path.startswith("/gerar_word_relatorio_hibrido"):
                ctx = payload.get("context", {})
                map_png_base64 = payload.get("map_png_base64") or None
                filename_code = re.sub(r"[^A-Za-z0-9_-]+", "_", str(ctx.get("codigo") or "setor")).strip("_") or "setor"
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"relatorio_hibrido_ia_{filename_code}_{timestamp}.docx"
                docx_bytes, engine, model = hybrid_docx_bytes(ctx, map_png_base64)
                RELATORIOS_DIR.mkdir(parents=True, exist_ok=True)
                output_path = RELATORIOS_DIR / filename
                output_path.write_bytes(docx_bytes)
                json_response(
                    self,
                    200,
                    {
                        "ok": True,
                        "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
                        "filename": filename,
                        "saved_path": str(output_path),
                        "engine": engine,
                        "model": model,
                    },
                )
                return
            json_response(self, 404, {"ok": False, "error": "Endpoint nao encontrado."})
        except Exception as exc:
            json_response(self, 500, {"ok": False, "error": str(exc)})

    def log_message(self, format: str, *args: Any) -> None:
        print(f"[IA contratos] {self.address_string()} - {format % args}")


def main() -> None:
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    print(f"Servico de IA para contratos em http://{HOST}:{PORT}")
    print("Opcional: defina OLLAMA_MODEL=nome-do-modelo para usar um modelo local do Ollama na redacao.")
    server.serve_forever()


if __name__ == "__main__":
    main()
